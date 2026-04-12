#!/usr/bin/env python3
"""Generate SiloController documentation as a Word .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ────────────────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def code(text):
    p = doc.add_paragraph(style="No Spacing")
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x40)
    p.paragraph_format.left_indent = Cm(1.0)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    return p

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# Title page
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SILO CONTROLLER")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x6c)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Patrick Blackett — Silo 2 Controller")
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Technical Reference & User Manual")
run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(datetime.date.today().strftime("%B %Y"))
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1 — Introduction
# ══════════════════════════════════════════════════════════════════════════════

heading("Introduction")
para(
    "The Silo Controller is a Raspberry Pi 5 application designed to control a "
    "linear actuator that opens and closes a silo lid fitted to a vessel. It "
    "simultaneously records navigation and attitude telemetry from an ArduPilot "
    "autopilot via MAVLink, records video from a Pi Camera with a live telemetry "
    "overlay, and provides control via a web browser, MQTT broker, and HTTP REST API."
)
para(
    "The system is designed to run headlessly under systemd with all operator "
    "interaction through a browser at port 5000. An optional Tk desktop GUI is "
    "available when a display is connected."
)

heading("Scope", level=2)
bullet("Control the silo lid (open/close) via relay-driven linear actuator")
bullet("Record MAVLink telemetry to CSV files")
bullet("Stream and record Pi Camera video with telemetry overlay burned in")
bullet("Accept commands and publish state via MQTT")
bullet("Provide a full REST HTTP API")
bullet("Configurable travel time, persistent relay state, and stress-test facility")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2 — Hardware
# ══════════════════════════════════════════════════════════════════════════════

heading("Hardware")

heading("Bill of Materials", level=2)
add_table(
    ["Item", "Part", "Notes"],
    [
        ["Compute", "Raspberry Pi 5", "Main controller"],
        ["Relay", "Eletechsup DR25E01", "3 A, 1-ch, bistable DPDT, 5 V logic"],
        ["Actuator", "Linak LA 121P00-1101220", "12 V DC, internal end-stop switches"],
        ["Autopilot/FC", "Matek H743", "ArduPilot, USB serial to Pi"],
        ["GPS", "u-blox M10Q", "Via ArduPilot (GPS1_TYPE=9). Supercapacitor backed."],
        ["RF Link", "mLRS", "MAVLink over radio"],
        ["Camera", "Pi Camera Module", "Optional — auto-detected"],
        ["Power (motor)", "12 V PSU", "Separate from Pi logic supply"],
    ],
    widths=[3.5, 5.0, 7.0],
)

heading("GPIO Wiring (BCM numbering)", level=2)
para(
    "The DR25E01 bistable relay uses a single trigger input (T). It toggles "
    "to the opposite position on each LOW pulse and latches without power."
)
add_table(
    ["Pi Pin", "DR25E01 Pin", "Function"],
    [
        ["GPIO 17", "T  (Trigger)", "100 ms LOW pulse toggles relay"],
        ["5 V pin", "V  (Coil power)", "Held HIGH at idle"],
        ["GND", "G  (Ground)", ""],
    ],
    widths=[3.5, 4.5, 7.5],
)
para(
    "Relay OUT+/OUT− connect to the actuator motor terminals. "
    "If open/close direction is reversed, swap OUT+/OUT−. "
    "A separate 12 V supply powers the actuator through the relay."
)

heading("How the Bistable Relay Works", level=2)
para(
    "The relay has no position feedback output. The software tracks which "
    "position the relay is in (relay_open bool) and only pulses pin T when "
    "a direction change is required. This state is persisted to silo_state.json "
    "and survives reboots. If the physical relay ever gets out of sync with the "
    "software state (e.g. after a power cut mid-travel), use the Configuration "
    "tab in the web UI to declare the actual position."
)
code("Idle:    GPIO 17 HIGH  — relay holds its last position (no power needed)\n"
     "Toggle:  GPIO 17 LOW for 100 ms  — relay flips, then GPIO 17 returns HIGH")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3 — Software Architecture
# ══════════════════════════════════════════════════════════════════════════════

heading("Software Architecture")
para(
    "The application is written in Python 3 and structured as a set of "
    "independent modules, each running on daemon threads. main.py is the "
    "single entry point."
)

heading("Module Overview", level=2)
add_table(
    ["Module", "File", "Role"],
    [
        ["SiloController", "core/silo.py", "State machine + GPIO driver. Thread-safe open/close/declare."],
        ["TelemetryReader", "telemetry/reader.py", "MAVLink reader; auto-detects serial port; provides live frame + diagnostics."],
        ["Recorder", "telemetry/recorder.py", "CSV telemetry logger. One file per session. PB*.csv naming."],
        ["StatsTracker", "telemetry/stats.py", "Rolling 1/5/10/30 min min/avg/max statistics."],
        ["CameraRecorder", "camera/recorder.py", "Pi camera auto-detect, H264 MP4 recording, MJPEG stream, telem overlay."],
        ["Flask app", "web/app.py", "REST API + SSE at port 5000."],
        ["MqttHandler", "mqtt/handler.py", "Paho MQTT client. Commands in, state out, events out."],
        ["SiloGUI", "gui/app.py", "Optional Tk desktop GUI (skipped if no display)."],
    ],
    widths=[4.0, 4.5, 7.0],
)

heading("Thread Model", level=2)
bullet("main thread — Tk event loop (or signal.pause() headless)")
bullet("mavlink-reader — serial MAVLink reader loop")
bullet("web-server — Flask threaded server")
bullet("sse-telem — pushes telemetry to SSE clients at 5 Hz")
bullet("sse-hb — heartbeat every 5 s")
bullet("actuator-open / actuator-close — short-lived per move")
bullet("stress-test — long-lived during stress test only")
bullet("cam-capture — continuous MJPEG frame capture (if camera present)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4 — Web Interface
# ══════════════════════════════════════════════════════════════════════════════

heading("Web Interface")
para("Navigate to http://<Pi-IP>:5000 in any browser on the same network.")

heading("Tabs", level=2)

heading("Live", level=3)
para(
    "Shows the current silo state (open/closed/opening/closing with animated dot), "
    "recording state, GPS diagnostics (fix type, satellites, HDOP, VDOP, H/V accuracy), "
    "mLRS link quality, navigation data (lat, lon, alt, speed, COG, pitch, roll), "
    "and a motion analysis table showing rolling statistics."
)

heading("Video", level=3)
para(
    "The second tab. Shows the live MJPEG camera feed with telemetry overlay burned "
    "into the image (time, speed, pitch, roll, COG, GPS fix). Telemetry data is also "
    "displayed in panels to the right of the feed. Features:"
)
bullet("Military test-range aesthetic — dark background, gold borders, scanlines")
bullet("● RECORDING / ○ LIVE indicator")
bullet("Snapshot button — downloads JPEG immediately")
bullet("Manual record start/stop")
bullet("Custom status message input — text is burned into video overlay and logged")
bullet("Saved video list with MP4 download")
bullet("Telem panel: GPS fix, lat/lon/alt, satellites, speed, COG, pitch, roll, silo state")

heading("History", level=3)
para(
    "Interactive track map (English Channel coastline), and scrollable/zoomable "
    "time-series charts for speed, heading, pitch, roll, pitch/roll/yaw rates, and altitude."
)

heading("Diagnostics", level=3)
para("Full autopilot diagnostics including:")
bullet("Autopilot firmware version, board UID, capabilities")
bullet("System status, EKF health flags and variances")
bullet("GPS signal detail (fix type, satellites, HDOP, VDOP, accuracy)")
bullet("GPS module hardware (type, HW/SW version if captured at boot)")
bullet("Satellite sky plot and per-satellite SNR table")
bullet("mLRS radio link quality, RSSI, buffer, errors")
bullet("Battery voltage, current, remaining %")
bullet("FC status log (STATUSTEXT messages from ArduPilot)")

heading("Log Files", level=3)
para(
    "Lists all CSV recording files with size, row count, and date. Clicking a "
    "file shows a full summary: start/end time, duration, motion statistics "
    "(min/avg/max for speed, pitch, roll, rates, altitude), GPS bounds, silo "
    "events, GPS fix quality histogram, and a mini track map. Direct CSV download."
)

heading("Configuration", level=3)
bullet("Declare Silo Position — manually declare open or closed without moving the actuator")
bullet("Actuator Travel Time — set the software timer (0.5–120 s) that matches actual stroke time")
bullet("Stress Test — run N open/close cycles with configurable pause, live progress bar, stop button")
bullet("Relay Wiring reference — GPIO 17, pulse width, state file location")

heading("API Docs", level=3)
para(
    "Shows exact curl and mosquitto commands with the Pi's real IP address "
    "pre-filled. Every code block has a one-click Copy button. Covers silo "
    "control, recording, telemetry queries, SSE, MQTT topics, camera message API, "
    "and GPIO wiring."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5 — REST API
# ══════════════════════════════════════════════════════════════════════════════

heading("REST API")
para("Base URL: http://<Pi-IP>:5000   No authentication required.")

heading("Silo Control", level=2)
add_table(
    ["Method", "Path", "Description", "Response"],
    [
        ["POST", "/api/open", "Open silo lid", '{"ok": true}'],
        ["POST", "/api/close", "Close silo lid", '{"ok": true}'],
        ["GET", "/api/status", "State + recording + uptime", '{"state":"open","recording":false,"uptime_s":3612}'],
        ["POST", "/api/config/declare", 'Declare position {"open":true}', '{"ok":true,"state":"open"}'],
        ["GET", "/api/config/relay_state", "Relay open bool + state + travel_time", "JSON"],
        ["GET", "/api/config/travel_time", "Current travel time (s)", '{"travel_time":10.0}'],
        ["POST", "/api/config/travel_time", '{"travel_time":12.5}', '{"ok":true,"travel_time":12.5}'],
    ],
    widths=[1.8, 4.5, 4.5, 4.7],
)

heading("Telemetry", level=2)
add_table(
    ["Method", "Path", "Description"],
    [
        ["GET", "/api/telemetry", "Live telemetry frame (lat, lon, alt, speed, pitch, roll, GPS fix, etc.)"],
        ["GET", "/api/diagnostics", "Full diagnostics dict including GPS frame fields, EKF, radio, battery"],
        ["GET", "/api/events", "SSE stream: state, telemetry, recording, cam_recording, heartbeat"],
        ["GET", "/api/history", "Historical telemetry rows for charts"],
        ["GET", "/api/track", "GPS track points for the map"],
    ],
    widths=[1.8, 4.5, 9.2],
)

heading("Camera", level=2)
add_table(
    ["Method", "Path", "Description"],
    [
        ["GET", "/api/camera/status", "Camera detected + recording state"],
        ["GET", "/api/camera/stream", "MJPEG live stream (open in browser or VLC)"],
        ["GET", "/api/camera/snapshot", "JPEG download of current frame"],
        ["POST", "/api/camera/message", '{"message":"text"} — add overlay annotation'],
        ["GET", "/api/camera/messages", "All messages log"],
        ["GET", "/api/camera/videos", "List saved MP4 files"],
        ["GET", "/api/camera/videos/<name>/download", "Download a specific MP4"],
        ["POST", "/api/camera/record/start", "Manual recording start"],
        ["POST", "/api/camera/record/stop", "Manual recording stop"],
    ],
    widths=[1.8, 5.2, 8.5],
)

heading("Stress Test", level=2)
add_table(
    ["Method", "Path", "Description"],
    [
        ["POST", "/api/stress/start", '{"cycles":10,"pause_s":3} — start automated cycle test'],
        ["POST", "/api/stress/stop", "Abort stress test"],
        ["GET", "/api/stress/status", "running, done, total, step"],
    ],
    widths=[1.8, 4.5, 9.2],
)

heading("Example curl Commands", level=2)
code("# Open the silo\ncurl -X POST http://192.168.1.100:5000/api/open\n")
code("# Check status\ncurl http://192.168.1.100:5000/api/status\n")
code("# Post a camera overlay message\ncurl -X POST http://192.168.1.100:5000/api/camera/message \\\n"
     "  -H 'Content-Type: application/json' \\\n"
     "  -d '{\"message\":\"LAUNCH SEQUENCE ARMED\"}'\n")
code("# Download latest video\ncurl -O http://192.168.1.100:5000/api/camera/videos/PBvideo_20260412_031100.mp4/download\n")
code("# Run 5-cycle stress test with 5 s pause\ncurl -X POST http://192.168.1.100:5000/api/stress/start \\\n"
     "  -H 'Content-Type: application/json' \\\n"
     "  -d '{\"cycles\":5,\"pause_s\":5}'\n")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6 — MQTT
# ══════════════════════════════════════════════════════════════════════════════

heading("MQTT Interface")
para(
    "A Mosquitto broker runs on the Pi at localhost:1883 with no authentication. "
    "The broker is accessible from any device on the same network."
)

heading("Installing the Broker", level=2)
code("sudo apt install mosquitto mosquitto-clients\n"
     "sudo systemctl enable --now mosquitto")

heading("Topics", level=2)
add_table(
    ["Direction", "Topic", "Payload", "Notes"],
    [
        ["Subscribe", "silo/command", "OPEN / CLOSE / STATUS", "Any case accepted"],
        ["Subscribe", "silo/record", "START / STOP / STATUS", "CSV telemetry recording"],
        ["Subscribe", "silo/message", "Any text", "Posted as camera overlay message"],
        ["Publish (retained)", "silo/status", "open / closed / opening / closing", "Immediate on subscribe"],
        ["Publish (retained)", "silo/recording", "true / false", "Immediate on subscribe"],
        ["Publish", "silo/event", 'JSON {"event":…,"ts":…,"source":…}', "Every state transition"],
    ],
    widths=[3.2, 3.5, 4.8, 4.0],
)

heading("Example mosquitto Commands", level=2)
code("# Open silo\nmosquitto_pub -h 192.168.1.100 -t silo/command -m OPEN\n")
code("# Close silo\nmosquitto_pub -h 192.168.1.100 -t silo/command -m CLOSE\n")
code("# Post a status message (overlay on video)\n"
     'mosquitto_pub -h 192.168.1.100 -t silo/message -m "TORPEDO DOOR OPEN"\n')
code("# Monitor silo state changes\nmosquitto_sub -h 192.168.1.100 -t silo/status\n")
code("# Subscribe to all silo events with timestamps\nmosquitto_sub -h 192.168.1.100 -t silo/event\n")
code("# Example event payload:\n"
     '{"event": "opening", "ts": "2026-04-12T03:11:00Z", "source": "web"}')

heading("Silo Event Source Values", level=2)
add_table(
    ["Source", "Triggered by"],
    [
        ["web", "Browser UI button"],
        ["mqtt", "MQTT silo/command"],
        ["api", "HTTP REST call"],
        ["config-ui", "Configuration tab declare button"],
        ["stress-test", "Automated stress test"],
        ["actuator", "Internal — final open/closed state after travel"],
        ["manual", "Configuration tab declare"],
    ],
    widths=[4.0, 11.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7 — Camera & Video
# ══════════════════════════════════════════════════════════════════════════════

heading("Camera and Video")

heading("Auto-Detection", level=2)
para(
    "On startup the application calls Picamera2.global_camera_info(). If one or "
    "more cameras are detected the CameraRecorder initialises immediately. If no "
    "camera is found, all camera API endpoints return gracefully (204 No Content "
    "for stream/snapshot, JSON with camera_detected:false for status) and all "
    "other application functions continue normally."
)

heading("Automatic Recording", level=2)
para(
    "Recording starts automatically when the silo begins opening (state = opening) "
    "and stops 2 seconds after the silo reaches the closed state. This ensures the "
    "full open-to-close sequence is captured including the final moments of closing."
)
para(
    "Recording can also be started and stopped manually via the Video tab, "
    "the REST API, or MQTT."
)

heading("Video Overlay", level=2)
para(
    "A telemetry overlay is burned into both the recorded MP4 files and the live "
    "MJPEG stream. The overlay always shows:"
)
bullet("● REC (recording) or ○ LIVE (stream only) indicator")
bullet("Current UTC date and time")
bullet("Ground speed (knots)")
bullet("Altitude (metres)")
bullet("Pitch angle (degrees, +nose-up)")
bullet("Roll angle (degrees, +starboard)")
bullet("Course over ground (degrees)")
bullet("GPS fix type and satellite count")
bullet("Last 3 custom status messages with timestamps")

heading("Status Messages", level=2)
para(
    "Custom text messages can be sent via HTTP POST, MQTT publish, or the Video "
    "tab input field. Messages are timestamped (UTC), stored in a rolling log "
    "(last 20 kept), and the most recent 3 are shown on the video overlay. "
    "They persist in memory until the service restarts."
)
para(
    "System events are automatically posted as messages: when the silo begins "
    "opening, reaches open, begins closing, or reaches closed, a timestamped "
    "entry is added — e.g. 'SILO OPENING  [03:11:00 UTC]  via web'."
)

heading("Disk Management", level=2)
para(
    "Before starting each new recording, the application checks available disk "
    "space. If free space is below 15% of the total, the oldest MP4 file in "
    "recordings/video/ is deleted. This ensures the system never fills the SD card."
)

heading("Video File Format", level=2)
add_table(
    ["Parameter", "Value"],
    [
        ["Container", "MP4"],
        ["Video codec", "H.264"],
        ["Resolution", "1280 × 720"],
        ["Frame rate", "30 fps"],
        ["Bitrate", "4 Mbps"],
        ["Naming", "PBvideo_YYYYMMDD_HHMMSS.mp4"],
        ["Location", "recordings/video/"],
    ],
    widths=[5.0, 10.5],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8 — Silo State Machine
# ══════════════════════════════════════════════════════════════════════════════

heading("Silo State Machine")

heading("States", level=2)
add_table(
    ["State", "Meaning", "Relay", "Actuator"],
    [
        ["closed", "Lid fully closed (default)", "Retracted position", "Stationary"],
        ["opening", "Lid moving open", "Extended position", "Extending"],
        ["open", "Lid fully open", "Extended position", "Stationary (limit switch)"],
        ["closing", "Lid moving closed", "Retracted position", "Retracting"],
    ],
    widths=[2.5, 4.5, 3.5, 5.0],
)

heading("Transitions", level=2)
code("closed  →  opening  →  open\nopen    →  closing  →  closed")
para(
    "An open command when already open (or opening) is rejected with "
    '{"ok": false, "reason": "Already open"}. Similarly for close.'
)

heading("Travel Time", level=2)
para(
    "The travel_time setting (default 2 s, configurable 0.5–120 s) controls "
    "how long the software waits in the opening/closing state before declaring "
    "the move complete. The Linak actuator's internal limit switches stop the "
    "motor at both ends automatically — travel_time is a software display timer, "
    "not a hard motor cutoff. Set it slightly longer than the actual full-stroke time."
)

heading("Stress Test", level=2)
para(
    "The stress test runs N open/close cycles with a configurable pause between "
    "each move. It waits for actual state transitions (polling silo.state) rather "
    "than sleeping — so it correctly handles any travel time setting. The test can "
    "be stopped at any point; the silo will complete its current move before stopping."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9 — Telemetry Recording
# ══════════════════════════════════════════════════════════════════════════════

heading("Telemetry Recording")
para(
    "CSV files are written to recordings/ with the naming convention "
    "PB_YYYYMMDD_HHMMSS.csv. Each file covers one recording session "
    "(start/stop via UI, MQTT, or API)."
)

heading("CSV Columns", level=2)
para("Key fields written per row:")
bullet("timestamp_unix — Unix timestamp (float)")
bullet("timestamp_iso — ISO 8601 UTC string")
bullet("lat, lon — decimal degrees")
bullet("alt_m — altitude metres MSL")
bullet("heading_deg — course over ground degrees")
bullet("groundspeed_ms — speed m/s")
bullet("pitch_deg, roll_deg, yaw_deg — attitude degrees")
bullet("pitch_rate_deg_s, roll_rate_deg_s, yaw_rate_deg_s — angular rates")
bullet("gps_fix — 0=none, 2=2D, 3=3D, 4=DGPS, 5=RTK float, 6=RTK fixed")
bullet("satellites — satellite count")
bullet("hdop, vdop — dilution of precision")
bullet("silo_state, silo_event, silo_source — silo transitions when they occur")

heading("GPS Assistance (Cold Start)", level=2)
para(
    "The script tools/gps_assist.py injects cold-start assistance data directly "
    "to the u-blox GPS via MAVLink GPS_INJECT_DATA messages, without any internet "
    "registration. This can reduce GPS acquisition time from ~15 minutes to 30–90 "
    "seconds outdoors."
)
para("Three types of data are injected:")
bullet("Current UTC time (from Pi NTP-synced clock)")
bullet("Approximate position (default: PO18 9AB, Bosham — configurable)")
bullet("Current GPS almanac (YUMA format from NAVCEN USCG — free, no login)")
code("python tools/gps_assist.py\n"
     "python tools/gps_assist.py --lat 50.83 --lon -0.87 --port /dev/ttyACM1")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10 — Installation & Service
# ══════════════════════════════════════════════════════════════════════════════

heading("Installation")

heading("Dependencies", level=2)
code("# Python packages\npip3 install flask paho-mqtt gpiozero pymavlink\n"
     "pip3 install picamera2 opencv-python  # optional — for camera\n\n"
     "# System packages\nsudo apt install mosquitto mosquitto-clients ffmpeg")

heading("Running Manually", level=2)
code("cd /home/pi/SiloController\npython3 main.py")

heading("Running as a systemd Service", level=2)
code("[Unit]\nDescription=Silo Controller\nAfter=network.target\n\n"
     "[Service]\nExecStart=/usr/bin/python3 /home/pi/SiloController/main.py\n"
     "WorkingDirectory=/home/pi/SiloController\nRestart=always\nRestartSec=3\nUser=pi\n\n"
     "[Install]\nWantedBy=multi-user.target")
code("sudo systemctl enable silo\nsudo systemctl start silo\nsudo journalctl -u silo -f")

heading("Checking Status", level=2)
code("sudo systemctl status silo        # service health\n"
     "sudo journalctl -u silo -n 50     # last 50 log lines\n"
     "curl http://localhost:5000/api/status  # API health check")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11 — Troubleshooting
# ══════════════════════════════════════════════════════════════════════════════

heading("Troubleshooting")
add_table(
    ["Symptom", "Likely Cause", "Fix"],
    [
        ["Silo opens when it should close (or vice versa)",
         "Relay OUT+/OUT− wired backwards",
         "Swap OUT+ and OUT− on the actuator motor terminals"],
        ["Software state out of sync with physical silo",
         "Power cut during travel, or manual override",
         "Use Configuration tab → Declare Silo Position"],
        ["No GPS fix",
         "Cold start, blocked antenna, or supercapacitor depleted",
         "Run tools/gps_assist.py to inject almanac and time assistance"],
        ["GPS fix/satellite count not shown in Diagnostics",
         "Known: GPS frame data not in diagnostics dict before this fix",
         "Ensure web/app.py merges frame data into /api/diagnostics response"],
        ["MAVLink device not found",
         "USB serial not detected, or wrong port",
         "Check /dev/ttyACM* exists; plug in Matek via USB; check baud rate (57600)"],
        ["Camera not detected",
         "picamera2 not installed, or camera cable not seated",
         "pip3 install picamera2; check ribbon cable; run Picamera2.global_camera_info()"],
        ["MQTT commands not working",
         "Mosquitto broker not running",
         "sudo systemctl start mosquitto"],
        ["Video overlay not visible",
         "opencv-python not installed",
         "pip3 install opencv-python"],
        ["Disk full — video not recording",
         "SD card full (< 15% free check triggered delete)",
         "Delete old recordings manually, or increase SD card size"],
    ],
    widths=[4.5, 5.0, 6.0],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12 — File Structure
# ══════════════════════════════════════════════════════════════════════════════

heading("File Structure")
code(
    "SiloController/\n"
    "├── main.py                    Entry point\n"
    "├── CLAUDE.md                  AI context file for Claude Code\n"
    "├── requirements.txt           Python dependencies\n"
    "├── silo_state.json            Persistent relay state (auto-generated)\n"
    "│\n"
    "├── core/\n"
    "│   └── silo.py                SiloController — state machine + GPIO\n"
    "│\n"
    "├── telemetry/\n"
    "│   ├── reader.py              MAVLink reader + TelemetryFrame\n"
    "│   ├── recorder.py            CSV recording\n"
    "│   └── stats.py               Rolling statistics\n"
    "│\n"
    "├── camera/\n"
    "│   └── recorder.py            Pi camera + MJPEG + MP4 + overlay\n"
    "│\n"
    "├── web/\n"
    "│   ├── app.py                 Flask REST API + SSE\n"
    "│   └── templates/\n"
    "│       └── index.html         Single-page browser UI\n"
    "│\n"
    "├── mqtt/\n"
    "│   └── handler.py             Paho MQTT client\n"
    "│\n"
    "├── gui/\n"
    "│   └── app.py                 Optional Tk desktop GUI\n"
    "│\n"
    "├── tools/\n"
    "│   └── gps_assist.py          GPS cold-start assistance injector\n"
    "│\n"
    "└── recordings/\n"
    "    ├── PB*.csv                CSV telemetry files\n"
    "    └── video/\n"
    "        └── PBvideo_*.mp4      H264 video recordings\n"
)

# ── Save ───────────────────────────────────────────────────────────────────────
out = "/home/pi/SiloController/SiloController_Documentation.docx"
doc.save(out)
print(f"Saved: {out}")
